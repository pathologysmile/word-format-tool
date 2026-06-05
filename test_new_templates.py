"""测试工作汇报和会议纪要格式模板"""

import os
import sys
from docx import Document
from docx.shared import Pt, Cm
from src.document_formatter import DocumentFormatter
from src.title_detector import TitleDetector
from src.document_parser import DocumentParser
import io


def create_test_work_report():
    """创建测试工作汇报文档"""
    doc = Document()
    
    # 添加标题
    doc.add_paragraph("2024年度工作总结报告", style='Title')
    
    # 一级标题
    p1 = doc.add_paragraph("一、工作回顾")
    p1.style.font.size = Pt(16)
    
    # 二级标题
    p2 = doc.add_paragraph("（一）核心业绩")
    
    # 正文
    p3 = doc.add_paragraph("    本年度完成了多个重要项目，取得了显著成效。在项目推进过程中，团队协作紧密，克服重重困难，最终实现了预期目标。")
    
    # 三级标题
    p4 = doc.add_paragraph("1. 重点项目A")
    
    # 正文
    p5 = doc.add_paragraph("    项目A是本年度的核心工作之一，涉及多个部门的协同配合。通过科学规划和有效执行，项目按期完成并取得了良好的经济效益。")
    
    # 一级标题
    p6 = doc.add_paragraph("二、存在问题")
    
    # 正文
    p7 = doc.add_paragraph("    在工作中也发现了一些问题，主要包括资源配置不够合理、沟通效率有待提高等。这些问题需要在下一年度重点解决。")
    
    # 一级标题
    p8 = doc.add_paragraph("三、下一步计划")
    
    # 正文
    p9 = doc.add_paragraph("    下一年度将继续深化改革，优化工作流程，提升工作效率。重点推进以下几方面工作：加强团队建设、完善制度体系、推进数字化转型。")
    
    # 保存到临时文件
    temp_path = "tests/sample_docs/test_work_report.docx"
    os.makedirs(os.path.dirname(temp_path), exist_ok=True)
    doc.save(temp_path)
    return temp_path


def create_test_meeting_minutes():
    """创建测试会议纪要文档"""
    doc = Document()
    
    # 添加标题
    doc.add_paragraph("XX项目周例会纪要", style='Title')
    
    # 会议信息
    p1 = doc.add_paragraph("会议主题：2024年Q3产品规划评审会")
    p2 = doc.add_paragraph("会议时间：2024-07-15 14:00 - 16:00")
    p3 = doc.add_paragraph("会议地点：301会议室")
    p4 = doc.add_paragraph("主持人：张三")
    p5 = doc.add_paragraph("记录人：李四")
    p6 = doc.add_paragraph("参会人员：王五、赵六、孙七（请假：周八）")
    
    # 一级标题
    p7 = doc.add_paragraph("一、会议决议")
    
    # 正文
    p8 = doc.add_paragraph("1. 通过《Q3产品路线图》，预算调整为50万。")
    p9 = doc.add_paragraph("2. 确定新产品上线时间为8月15日。")
    
    # 一级标题
    p10 = doc.add_paragraph("二、主要讨论内容")
    
    # 二级标题
    p11 = doc.add_paragraph("（一）关于用户增长放缓的问题")
    
    # 三级标题
    p12 = doc.add_paragraph("1. 现状：最近两个月用户增长率下降了15%")
    p13 = doc.add_paragraph("2. 原因：市场竞争加剧，产品功能更新不及时")
    
    # 一级标题
    p14 = doc.add_paragraph("三、待办事项")
    
    # 创建待办事项表格
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    
    # 表头
    headers = ["序号", "任务内容", "责任人", "截止日期"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    
    # 表格内容
    data = [
        ["1", "完成用户调研报告", "王五", "2024-07-25"],
        ["2", "优化产品功能清单", "赵六", "2024-07-30"],
        ["3", "制定营销推广方案", "孙七", "2024-08-05"]
    ]
    
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = cell_data
    
    # 保存到临时文件
    temp_path = "tests/sample_docs/test_meeting_minutes.docx"
    os.makedirs(os.path.dirname(temp_path), exist_ok=True)
    doc.save(temp_path)
    return temp_path


def test_work_report_formatting():
    """测试工作汇报格式化"""
    print("=" * 60)
    print("测试工作汇报格式转换")
    print("=" * 60)
    
    # 创建测试文档
    test_file = create_test_work_report()
    print(f"✅ 测试文档已创建: {test_file}")
    
    # 格式化
    template_path = "templates/work_report.json"
    if not os.path.exists(template_path):
        print(f"❌ 模板文件不存在: {template_path}")
        return
    
    output_path = "tests/sample_docs/test_work_report_formatted.docx"
    
    try:
        formatter = DocumentFormatter(template_path)
        
        # 解析文档
        with open(test_file, 'rb') as f:
            file_stream = io.BytesIO(f.read())
        
        parser = DocumentParser(file_stream)
        paragraphs = parser.extract_paragraphs()
        
        # 识别标题
        detector = TitleDetector(paragraphs)
        detected_titles = detector.hybrid_detect()
        
        # 应用格式化
        output_stream = io.BytesIO()
        file_stream.seek(0)
        formatter.apply_formatting(file_stream, output_stream, detected_levels=detected_titles)
        
        # 保存结果
        output_stream.seek(0)
        with open(output_path, 'wb') as f:
            f.write(output_stream.getvalue())
        
        print(f"✅ 格式化完成: {output_path}")
        
    except Exception as e:
        print(f"❌ 格式化失败: {str(e)}")
        import traceback
        traceback.print_exc()


def test_meeting_minutes_formatting():
    """测试会议纪要格式化"""
    print("\n" + "=" * 60)
    print("测试会议纪要格式转换")
    print("=" * 60)
    
    # 创建测试文档
    test_file = create_test_meeting_minutes()
    print(f"✅ 测试文档已创建: {test_file}")
    
    # 格式化
    template_path = "templates/meeting_minutes.json"
    if not os.path.exists(template_path):
        print(f"❌ 模板文件不存在: {template_path}")
        return
    
    output_path = "tests/sample_docs/test_meeting_minutes_formatted.docx"
    
    try:
        formatter = DocumentFormatter(template_path)
        
        # 解析文档
        with open(test_file, 'rb') as f:
            file_stream = io.BytesIO(f.read())
        
        parser = DocumentParser(file_stream)
        paragraphs = parser.extract_paragraphs()
        
        # 识别标题
        detector = TitleDetector(paragraphs)
        detected_titles = detector.hybrid_detect()
        
        # 应用格式化
        output_stream = io.BytesIO()
        file_stream.seek(0)
        formatter.apply_formatting(file_stream, output_stream, detected_levels=detected_titles)
        
        # 保存结果
        output_stream.seek(0)
        with open(output_path, 'wb') as f:
            f.write(output_stream.getvalue())
        
        print(f"✅ 格式化完成: {output_path}")
        
    except Exception as e:
        print(f"❌ 格式化失败: {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    # 运行测试
    test_work_report_formatting()
    test_meeting_minutes_formatting()
    
    print("\n" + "=" * 60)
    print("测试完成！")
    print("=" * 60)
