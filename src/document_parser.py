"""文档解析器 - 读取 .docx 文件并提取结构化信息"""

from docx import Document
from typing import List, Dict, Union
import io
import os


class DocumentParser:
    """Word 文档解析器"""
    
    def __init__(self, file_source: Union[str, io.BytesIO]):
        """
        初始化文档解析器
        
        Args:
            file_source: .docx 文件路径或 BytesIO 对象
        """
        self.file_source = file_source
        
        # 支持文件路径或内存流
        if isinstance(file_source, str):
            if not os.path.exists(file_source):
                raise FileNotFoundError(f"文件不存在: {file_source}")
            if not file_source.endswith('.docx'):
                raise ValueError("仅支持 .docx 格式文件")
            self.doc = Document(file_source)
            self.file_path = file_source
        elif isinstance(file_source, io.BytesIO):
            self.doc = Document(file_source)
            self.file_path = "<memory>"
        else:
            raise TypeError("不支持的文件类型，需要 str 或 BytesIO")
    
    def extract_paragraphs(self) -> List[Dict]:
        """
        提取所有段落信息
        
        Returns:
            List[Dict]: 段落信息列表
            [
                {
                    'index': 0,
                    'text': '段落文本',
                    'style': 'Normal',
                    'runs_count': 1
                },
                ...
            ]
        """
        paragraphs = []
        for i, para in enumerate(self.doc.paragraphs):
            paragraphs.append({
                'index': i,
                'text': para.text,
                'preview': para.text[:50] + ('...' if len(para.text) > 50 else ''),
                'style': para.style.name if para.style else 'Normal',
                'runs_count': len(para.runs),
                'is_empty': len(para.text.strip()) == 0
            })
        
        return paragraphs
    
    def extract_tables(self) -> List[Dict]:
        """
        提取表格信息
        
        Returns:
            List[Dict]: 表格信息列表
        """
        tables = []
        for i, table in enumerate(self.doc.tables):
            table_info = {
                'index': i,
                'rows': len(table.rows),
                'cols': len(table.columns),
                'cells': []
            }
            
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    table_info['cells'].append({
                        'row': row_idx,
                        'col': col_idx,
                        'text': cell.text
                    })
            
            tables.append(table_info)
        
        return tables
    
    def get_metadata(self) -> Dict:
        """
        获取文档元数据
        
        Returns:
            Dict: 元数据字典
        """
        metadata = {
            'paragraph_count': len(self.doc.paragraphs),
            'table_count': len(self.doc.tables),
            'section_count': len(self.doc.sections),
            'file_path': self.file_path,
        }
        
        # 只有文件路径时才计算文件大小
        if isinstance(self.file_source, str):
            metadata['file_size_kb'] = os.path.getsize(self.file_source) / 1024
        else:
            # 内存流估算大小
            self.file_source.seek(0, 2)  # 移动到末尾
            metadata['file_size_kb'] = self.file_source.tell() / 1024
            self.file_source.seek(0)  # 重置位置
        
        return metadata
    
    def get_page_setup(self) -> Dict:
        """
        获取当前页面设置
        
        Returns:
            Dict: 页面设置信息
        """
        section = self.doc.sections[0]  # 获取第一节
        
        return {
            'top_margin': section.top_margin.mm,
            'bottom_margin': section.bottom_margin.mm,
            'left_margin': section.left_margin.mm,
            'right_margin': section.right_margin.mm,
            'page_width': section.page_width.mm,
            'page_height': section.page_height.mm,
            'header_distance': section.header_distance.mm,
            'footer_distance': section.footer_distance.mm
        }
