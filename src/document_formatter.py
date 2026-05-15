"""文档格式化引擎 - 将模板应用到 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
import json
import io
import os
from typing import Dict, List, Union


class DocumentFormatter:
    """文档格式化器"""

    def __init__(self, template_path: str):
        """
        初始化格式化器
        
        Args:
            template_path: JSON 模板文件路径
        """
        with open(template_path, 'r', encoding='utf-8') as f:
            self.template = json.load(f)
        
    def apply_formatting(self, input_source: Union[str, io.BytesIO], output_dest: Union[str, io.BytesIO], detected_levels: List[Dict] = None):
        """
        应用模板格式化文档
        
        Args:
            input_source: 输入文档路径或 BytesIO 对象
            output_dest: 输出文档路径或 BytesIO 对象
            detected_levels: 标题识别结果列表 (来自 TitleDetector)
        """
        # 加载文档(支持文件路径或内存流)
        if isinstance(input_source, str):
            doc = Document(input_source)
        elif isinstance(input_source, io.BytesIO):
            input_source.seek(0)  # 确保在开头
            doc = Document(input_source)
        else:
            raise TypeError("input_source 必须是 str 或 BytesIO")
        
        # 1. 应用页面设置
        self._apply_page_setup(doc)
        
        # 2. 应用页眉页脚
        self._apply_header_footer(doc)
        
        # 3. 逐段落应用格式
        for i, para in enumerate(doc.paragraphs):
            level = 'body'
            if detected_levels and i < len(detected_levels):
                level = detected_levels[i].get('level', 'body')
            self._format_paragraph(para, level)
            
        # 4. 格式化表格
        self._format_tables(doc)
            
        # 5. 保存文档(支持文件路径或内存流)
        if isinstance(output_dest, str):
            doc.save(output_dest)
        elif isinstance(output_dest, io.BytesIO):
            doc.save(output_dest)
        else:
            raise TypeError("output_dest 必须是 str 或 BytesIO")
        
    def _apply_page_setup(self, doc: Document):
        """应用页面设置（边距、纸张大小）"""
        page_config = self.template.get('page', {})
        if not page_config:
            return
            
        section = doc.sections[0]
        
        # 设置边距 (转换为 cm)
        if 'top_margin_mm' in page_config:
            section.top_margin = Cm(page_config['top_margin_mm'] / 10)
        if 'bottom_margin_mm' in page_config:
            section.bottom_margin = Cm(page_config['bottom_margin_mm'] / 10)
        if 'left_margin_mm' in page_config:
            section.left_margin = Cm(page_config['left_margin_mm'] / 10)
        if 'right_margin_mm' in page_config:
            section.right_margin = Cm(page_config['right_margin_mm'] / 10)
            
        # 设置纸张大小
        if 'page_width_mm' in page_config:
            section.page_width = Cm(page_config['page_width_mm'] / 10)
        if 'page_height_mm' in page_config:
            section.page_height = Cm(page_config['page_height_mm'] / 10)

    def _format_paragraph(self, para, level: str = 'body'):
        """格式化单个段落，根据层级应用不同样式"""
        level_config = self.template.get('levels', {}).get(level, self.template.get('levels', {}).get('body', {}))
        self._apply_level_style(para, level_config)

    def _apply_header_footer(self, doc: Document):
        """应用页眉页脚配置"""
        header_config = self.template.get('header', {})
        footer_config = self.template.get('footer', {})
        
        section = doc.sections[0]
        
        # 处理页眉
        if header_config:
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.text = header_config.get('text', '')
            
            # 设置页眉字体和对齐
            font_name = header_config.get('font_main', '宋体')
            font_size = header_config.get('font_size_pt', 9)
            for run in header_para.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
                r = run._element.rPr.rFonts
                r.set(qn('w:eastAsia'), font_name)
            
            align_map = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER, 'right': WD_ALIGN_PARAGRAPH.RIGHT}
            header_para.alignment = align_map.get(header_config.get('alignment', 'center'), WD_ALIGN_PARAGRAPH.CENTER)

        # 处理页脚（逻辑同页眉）
        if footer_config:
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.text = footer_config.get('text', '')
            # ... (此处省略重复的字体设置代码，实际开发中可提取为公共方法)

    def _format_tables(self, doc: Document):
        """格式化文档中的所有表格"""
        table_config = self.template.get('table', {})
        if not table_config:
            return
            
        for table in doc.tables:
            # 设置表格对齐
            align_type = table_config.get('alignment', 'center')
            if align_type == 'center':
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
            elif align_type == 'left':
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                
            # 遍历单元格设置字体和边框
            font_name = table_config.get('font_main', '宋体')
            font_size = table_config.get('font_size_pt', 10.5)
            
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
                            r = run._element.rPr.rFonts
                            r.set(qn('w:eastAsia'), font_name)
                        
                        # 单元格内段落对齐
                        cell_align = table_config.get('cell_alignment', 'center')
                        align_map = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER, 'right': WD_ALIGN_PARAGRAPH.RIGHT}
                        para.alignment = align_map.get(cell_align, WD_ALIGN_PARAGRAPH.CENTER)

    def _apply_level_style(self, para, config: Dict):
        """应用指定层级的样式配置"""
        # 1. 字体设置（支持中英文/数字区分）
        font_main = config.get('font_main', '宋体')
        font_number = config.get('font_number', 'Times New Roman')
        font_size = config.get('font_size_pt', 12)
        is_bold = config.get('bold', False)
        
        for run in para.runs:
            # 简单的文本分离逻辑：如果包含数字或英文，则应用 font_number
            text = run.text
            if any(c.isdigit() or (c.isalpha() and c.isascii()) for c in text):
                run.font.name = font_number
                r = run._element.rPr.rFonts
                r.set(qn('w:eastAsia'), font_main) # 中文部分保持主字体
                r.set(qn('w:ascii'), font_number)
                r.set(qn('w:hAnsi'), font_number)
            else:
                run.font.name = font_main
                r = run._element.rPr.rFonts
                r.set(qn('w:eastAsia'), font_main)
            
            run.font.size = Pt(font_size)
            run.font.bold = is_bold

        # 2. 对齐方式
        align_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        align_type = config.get('alignment', 'justify')
        para.alignment = align_map.get(align_type, WD_ALIGN_PARAGRAPH.JUSTIFY)

        # 3. 行距与间距（使用 XML 直接操作，确保 100% 生效）
        line_spacing_type = config.get('line_spacing_type', 'multiple')
        line_spacing_value = float(config.get('line_spacing_value', 1.0))
        
        # 兼容旧版模板的 'fixed' 写法
        if line_spacing_type == 'fixed':
            line_spacing_type = 'exactly'
        
        # 获取段落属性 XML
        pPr = para._element.pPr
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            para._element.insert(0, pPr)
        
        # 清除已有的行距设置
        existing = pPr.find(qn('w:spacing'))
        if existing is not None:
            pPr.remove(existing)
        
        # 创建新的 spacing 元素
        spacing = OxmlElement('w:spacing')
        
        if line_spacing_type == 'single':
            spacing.set(qn('w:lineRule'), 'auto')
            spacing.set(qn('w:line'), '240')  # 单倍行距默认值
        elif line_spacing_type == '1.5':
            spacing.set(qn('w:lineRule'), 'auto')
            spacing.set(qn('w:line'), '360')  # 1.5 * 240
        elif line_spacing_type == 'double':
            spacing.set(qn('w:lineRule'), 'auto')
            spacing.set(qn('w:line'), '480')  # 2.0 * 240
        elif line_spacing_type == 'at_least':
            spacing.set(qn('w:lineRule'), 'atLeast')
            spacing.set(qn('w:line'), str(int(line_spacing_value * 20)))  # Word 使用 1/20 磅为单位
        elif line_spacing_type == 'exactly':
            spacing.set(qn('w:lineRule'), 'exact')
            spacing.set(qn('w:line'), str(int(line_spacing_value * 20)))  # 固定值 30磅 = 600
        elif line_spacing_type == 'multiple':
            spacing.set(qn('w:lineRule'), 'auto')
            spacing.set(qn('w:line'), str(int(line_spacing_value * 240)))
            
        pPr.append(spacing)
            
        # 段前段后间距（支持按"行"计算）
        space_before_lines = float(config.get('space_before_lines', 0.0))
        space_after_lines = float(config.get('space_after_lines', 0.0))
        if space_before_lines > 0:
            para.space_before = Pt(line_spacing_value * space_before_lines)
        if space_after_lines > 0:
            para.space_after = Pt(line_spacing_value * space_after_lines)

        # 4. 首行缩进
        indent_chars = config.get('first_line_indent_chars', 0)
        if indent_chars > 0:
            # 假设一个字符约等于 0.74cm (基于五号字估算)，这里根据字号动态调整会更准
            para.first_line_indent = Cm(indent_chars * (font_size / 16) * 0.74)
