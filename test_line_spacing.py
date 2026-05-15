"""测试行距设置是否正确应用"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
import os

# 创建测试文档
doc = Document()

# 添加测试段落
para1 = doc.add_paragraph("这是测试段落1，用于验证固定行距30磅是否生效。")
para2 = doc.add_paragraph("这是测试段落2，如果行距设置正确，两行之间应该有明显的间距。")

# 尝试设置固定行距 30 磅
for para in [para1, para2]:
    # 方法1: 使用枚举常量
    para.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    para.line_spacing = Pt(30)
    
    # 设置字体
    for run in para.runs:
        run.font.name = '仿宋_GB2312'
        run.font.size = Pt(17)
        r = run._element.rPr.rFonts
        r.set(qn('w:eastAsia'), '仿宋_GB2312')

# 保存文档
output_path = 'temp/test_line_spacing.docx'
os.makedirs('temp', exist_ok=True)
doc.save(output_path)

print(f"✅ 测试文档已保存: {output_path}")
print("请打开该文档检查行距是否为 30 磅")

# 读取并验证
doc2 = Document(output_path)
for i, para in enumerate(doc2.paragraphs):
    print(f"\n段落 {i}:")
    print(f"  line_spacing_rule: {para.line_spacing_rule}")
    print(f"  line_spacing: {para.line_spacing}")
    print(f"  文本: {para.text[:50]}")
