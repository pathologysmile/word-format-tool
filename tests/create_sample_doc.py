"""创建测试用的 Word 文档"""
from docx import Document
from docx.shared import Pt


def create_sample_official_doc():
    """创建公文样例文档"""
    doc = Document()
    
    # 大标题
    title = doc.add_paragraph('关于测试的通知')
    title.style = 'Title'
    
    # 一级标题
    h1 = doc.add_paragraph('一、工作目标')
    h1.style = 'Heading 1'
    
    # 正文
    body = doc.add_paragraph('这是测试正文内容,用于验证格式化功能。')
    body.style = 'Normal'
    
    # 二级标题(无样式,靠正则识别)
    h2 = doc.add_paragraph('(二)具体措施')
    h2.style = 'Normal'
    
    # 更多正文
    body2 = doc.add_paragraph('具体的实施步骤和措施。')
    
    # 三级标题
    h3 = doc.add_paragraph('3.时间安排')
    h3.style = 'Normal'
    
    # 四级标题
    h4 = doc.add_paragraph('(4)注意事项')
    h4.style = 'Normal'
    
    doc.save('tests/sample_docs/sample_official.docx')
    print("✅ 测试文档创建成功: tests/sample_docs/sample_official.docx")


if __name__ == '__main__':
    import os
    os.makedirs('tests/sample_docs', exist_ok=True)
    create_sample_official_doc()
