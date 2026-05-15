"""深度测试行距设置 - 直接操作 XML"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_line_spacing_exact(para, value):
    """通过 XML 直接设置固定行距（最可靠的方法）"""
    pPr = para._element.pPr
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    
    # 清除已有的行距设置
    existing = pPr.find(qn('w:spacing'))
    if existing is not None:
        pPr.remove(existing)
    
    # 创建新的 spacing 元素
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:lineRule'), 'exact')  # exact 表示固定值
    spacing.set(qn('w:line'), str(int(value * 20)))  # Word 使用 1/20 磅为单位
    pPr.append(spacing)

# 创建测试文档
doc = Document()

para1 = doc.add_paragraph("测试段落1 - 这是用于验证固定行距30磅是否生效的测试文本。如果设置正确，行距应该是固定的30磅。")
para2 = doc.add_paragraph("测试段落2 - 第二段落用于对比。两行之间应该有明显的固定间距，而不是自动调整的行距。")

# 使用 XML 直接设置固定行距 30 磅
for para in [para1, para2]:
    set_line_spacing_exact(para, 30)
    
    # 设置字体
    for run in para.runs:
        run.font.name = '仿宋_GB2312'
        run.font.size = Pt(17)
        r = run._element.rPr.rFonts
        r.set(qn('w:eastAsia'), '仿宋_GB2312')

# 保存文档
output_path = 'temp/test_xml_line_spacing.docx'
os.makedirs('temp', exist_ok=True)
doc.save(output_path)

print(f"✅ XML 方法测试文档已保存: {output_path}")
print("请打开该文档检查行距是否为固定值 30 磅")

# 读取验证
doc2 = Document(output_path)
for i, para in enumerate(doc2.paragraphs):
    print(f"\n段落 {i}:")
    pPr = para._element.pPr
    if pPr is not None:
        spacing = pPr.find(qn('w:spacing'))
        if spacing is not None:
            print(f"  lineRule: {spacing.get(qn('w:lineRule'))}")
            print(f"  line: {spacing.get(qn('w:line'))}")
    print(f"  文本: {para.text[:30]}...")
